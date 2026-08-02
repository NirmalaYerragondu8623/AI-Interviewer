"""Parses an uploaded question-bank document into rows of
{topic, category, question_text, reference_answer}.

  - CSV / XLSX: a header row containing columns matching topic/category/question/answer.
  - DOCX: a table with the same four columns, if present; otherwise labeled paragraphs
    ("Topic: ...", "Question: ...", etc.).
  - PDF: tuned to the real "AI/ML & Data Science Interview Q&A Compendium" document —
    it has no tables; sections ("N. Section Name", bold ~17pt) each contain numbered
    questions ("Qn. ...", bold ~11.5pt) followed by an answer in regular/monospace text.
    Parsing uses font name + size to tell headers/questions/answers apart, since there's
    no other structural marker. Falls back to table extraction (the previous approach)
    if a differently-formatted PDF yields no font-based matches.

If a future document has a different layout, only the corresponding `_parse_*` function
below needs to change — `parse_question_bank_document` and everything downstream
(the ingest endpoint, the DB insert) stays the same.
"""

import csv
import io
import re

import pandas as pd

# Header text we'll accept for each logical column, matched case-insensitively
# with punctuation/whitespace stripped (e.g. "Reference Answer" -> "referenceanswer").
_COLUMN_ALIASES = {
    "topic": {"topic"},
    "category": {"category", "subcategory", "type"},
    "question_text": {"question", "questiontext", "question text"},
    "reference_answer": {"referenceanswer", "answer", "modelanswer", "expectedanswer", "sampleanswer"},
}


def _normalize_header(value: str) -> str:
    return re.sub(r"[^a-z0-9]", "", value.lower())


def _map_columns(headers: list[str]) -> dict[str, str]:
    """Maps our logical field names to whatever the document's actual header text is."""
    normalized = {_normalize_header(h): h for h in headers}
    mapping: dict[str, str] = {}
    for field, aliases in _COLUMN_ALIASES.items():
        for alias in aliases:
            if _normalize_header(alias) in normalized:
                mapping[field] = normalized[_normalize_header(alias)]
                break
    missing = set(_COLUMN_ALIASES) - set(mapping)
    if missing:
        raise ValueError(
            f"Could not find columns for: {', '.join(sorted(missing))}. "
            f"Found headers: {headers}"
        )
    return mapping


def _rows_from_dataframe(df: "pd.DataFrame") -> list[dict]:
    mapping = _map_columns([str(c) for c in df.columns])
    rows = []
    for _, record in df.iterrows():
        row = {field: str(record[col]).strip() for field, col in mapping.items()}
        if row["question_text"] and row["question_text"].lower() != "nan":
            rows.append(row)
    return rows


def _parse_csv(content: bytes) -> list[dict]:
    text = content.decode("utf-8-sig")
    reader = csv.DictReader(io.StringIO(text))
    mapping = _map_columns(reader.fieldnames or [])
    rows = []
    for record in reader:
        row = {field: (record.get(col) or "").strip() for field, col in mapping.items()}
        if row["question_text"]:
            rows.append(row)
    return rows


def _parse_xlsx(content: bytes) -> list[dict]:
    df = pd.read_excel(io.BytesIO(content), engine="openpyxl")
    return _rows_from_dataframe(df)


def _parse_docx(content: bytes) -> list[dict]:
    import docx

    document = docx.Document(io.BytesIO(content))

    if document.tables:
        rows: list[dict] = []
        for table in document.tables:
            table_rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if not table_rows:
                continue
            headers, data_rows = table_rows[0], table_rows[1:]
            mapping = _map_columns(headers)
            col_index = {field: headers.index(col) for field, col in mapping.items()}
            for data_row in data_rows:
                row = {field: data_row[idx].strip() for field, idx in col_index.items()}
                if row["question_text"]:
                    rows.append(row)
        if rows:
            return rows

    # Fallback: labeled paragraphs, e.g. "Topic: X", "Question: Y", one entry per blank-line block.
    rows = []
    current: dict[str, str] = {}
    label_pattern = re.compile(r"^\s*(topic|category|question|answer|reference answer)\s*:\s*(.*)$", re.IGNORECASE)
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = label_pattern.match(text)
        if match:
            label, value = match.group(1).lower(), match.group(2).strip()
            field = {
                "topic": "topic",
                "category": "category",
                "question": "question_text",
                "answer": "reference_answer",
                "reference answer": "reference_answer",
            }[label]
            current[field] = value
        elif not text and current:
            if current.get("question_text"):
                rows.append(current)
            current = {}
    if current.get("question_text"):
        rows.append(current)
    return rows


def _classify_pdf_word(fontname: str, size: float) -> str:
    """Maps a word's font to its structural role in the compendium's layout."""
    size = round(size, 1)
    if "Bold" in fontname and size >= 15:
        return "header"  # section title, e.g. "1. LLM & GenAI Fundamentals"
    if "Bold" in fontname and 10.5 <= size <= 12.5:
        return "question"  # "Qn. ..."
    if fontname.split("-")[0].split(",")[0] in ("Helvetica", "Arial") and size <= 8.5:
        return "footer"  # page footer, e.g. "... Compendium Page 12"
    return "answer"  # regular prose or monospace code


def _parse_pdf_by_font(content: bytes) -> list[dict]:
    import pdfplumber

    tokens = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for word in page.extract_words(extra_attrs=["fontname", "size"]):
                role = _classify_pdf_word(word["fontname"], word["size"])
                if role == "footer":
                    continue
                tokens.append({"text": word["text"], "role": role, "top": round(word["top"], 1)})

    # Group consecutive same-role words into blocks. Within a block, a new line
    # (word "top" jumps) is preserved as a real line break — this matters for code
    # answers — while words on the same line are just space-joined.
    blocks: list[dict] = []
    current: dict | None = None
    for token in tokens:
        if current and current["role"] == token["role"]:
            if abs(token["top"] - current["last_top"]) > 2:
                current["lines"].append(token["text"])
            else:
                current["lines"][-1] += " " + token["text"]
            current["last_top"] = token["top"]
        else:
            if current:
                blocks.append(current)
            current = {"role": token["role"], "lines": [token["text"]], "last_top": token["top"]}
    if current:
        blocks.append(current)

    rows: list[dict] = []
    topic = ""
    pending_question = ""
    for block in blocks:
        if block["role"] == "header":
            topic = re.sub(r"^\d+\.\s*", "", " ".join(block["lines"])).strip()
        elif block["role"] == "question":
            text = re.sub(r"^Q\d+\.\s*", "", " ".join(block["lines"])).strip()
            # Guards against stray same-font fragments (e.g. an arrow glyph mid-question
            # resetting the font run) being mistaken for a new question.
            if re.match(r"^Q\d+\.", " ".join(block["lines"])) or not pending_question:
                pending_question = text
            else:
                pending_question += " " + text
        elif block["role"] == "answer" and pending_question:
            answer = "\n".join(block["lines"]).strip()
            # A very short "answer" immediately after a question is almost certainly a
            # stray glyph (e.g. an arrow icon) breaking up the question's font run, not
            # a real answer — skip it and keep accumulating the question instead of
            # prematurely closing it out.
            if len(answer) < 8:
                continue
            if len(pending_question.split()) >= 3:
                rows.append(
                    {
                        "topic": topic,
                        "category": "",
                        "question_text": pending_question,
                        "reference_answer": answer,
                    }
                )
            pending_question = ""

    return rows


def _parse_pdf_by_table(content: bytes) -> list[dict]:
    import pdfplumber

    rows: list[dict] = []
    with pdfplumber.open(io.BytesIO(content)) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables():
                if not table or len(table) < 2:
                    continue
                headers, data_rows = table[0], table[1:]
                try:
                    mapping = _map_columns([h or "" for h in headers])
                except ValueError:
                    continue
                col_index = {field: headers.index(col) for field, col in mapping.items()}
                for data_row in data_rows:
                    row = {field: (data_row[idx] or "").strip() for field, idx in col_index.items()}
                    if row["question_text"]:
                        rows.append(row)
    return rows


def _parse_pdf(content: bytes) -> list[dict]:
    rows = _parse_pdf_by_font(content)
    if rows:
        return rows
    return _parse_pdf_by_table(content)


_PARSERS = {
    "csv": _parse_csv,
    "xlsx": _parse_xlsx,
    "xls": _parse_xlsx,
    "docx": _parse_docx,
    "pdf": _parse_pdf,
}


def parse_question_bank_document(filename: str, content: bytes) -> list[dict]:
    extension = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    parser = _PARSERS.get(extension)
    if parser is None:
        raise ValueError(f"Unsupported file type '.{extension}'. Supported: {', '.join(sorted(_PARSERS))}")

    rows = parser(content)
    if not rows:
        raise ValueError("No question rows found in the uploaded document.")
    return rows
